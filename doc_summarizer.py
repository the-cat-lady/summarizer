# Requirements:
# usage:  python doc_summarizer.py <input.pdf> <output.docx>

'''
apt-get install build-essential libpoppler-cpp-dev pkg-config python3-dev
#pip install docx
pip install python-docx
pip install pdftotext
pip install transformers'''

import argparse
import docx
import math

import pdftotext
from itertools import zip_longest

from transformers import pipeline
import os
import re

class Summary:

    '''NLP with BERT-HuggingFace Transformers!'''
    def __init__(self, infile_path, outfile_path, nlp="summarization",batch_size=2700, nlp_kwargs=None):
        print(f'Initializing {nlp} pipeline')
        self.nlp = pipeline(nlp)
        self.infile_path = infile_path
        self.text = ''

        print('Extracting Text')
        if self.infile_path[-3:] == 'pdf':
            self.pdf_get_text()

        elif self.infile_path[-3:] == 'vtt':
            self.vtt_get_text()

        self.ner = pipeline('ner', grouped_entities=True)
        self.outfile_path = outfile_path
        self.batch_size = batch_size
        self.summaries = []
        self.tags = set()

    def pdf_get_text(self):
        '''PDFs'''
        with open(self.infile_path, 'rb') as f:
            self.pages = pdftotext.PDF(f)
        self.text = '\n\n'.join(page for page in self.pages)

        print(self.pages[0])

    def vtt_get_text(self):  # DROP THIS??
        '''MS Stream Transcripts'''

        with open(self.infile_path, 'r') as f:
            transcript = f.read()
        keepers = []
        for line in transcript.split('\n')[1:]:
          if line == '' or 'NOTE' in line or '_' in line:
            pass
          else:
            keepers.append(line)

        self.text = ' '.join(keepers)

    def do_nlp(self):
        '''Summarize text scraped from links''' # NOT REALLY?
        N = len(self.text)
        print('\n\nInside do_nlp\n', N, self.text)
        # make sure n_batches is always at least 1
        n_batches = math.ceil((N+1) / self.batch_size)
        batch = N // n_batches # take the floor

        for i in range(0, N, batch):
            print(i, batch+i)
            section = self.text[i:(i+batch)]
            try:
                if len(section) < 50:
                    print('section too short')
                    continue

                summary = self.nlp(section, min_length=90, max_length=200)
                self.summaries.append(summary[0]['summary_text'])
                tag_set = set(x['word'] for x in self.ner(section))
                self.tags.update(tag_set)
                print(summary)
            except Exception as e:
                print(f'\nFAILURE: {e}\n\n')
                continue
        return self.summaries

    def clean_summaries(self):
        '''Clean summarized text'''
        print('Inside clean_summaries')
        self.final_text = '. '.join(sentence[0].upper() + sentence[1:] for sentence in '\n'.join(self.summaries).split(' . '))
        return self.final_text

    def create_text_section(self, title='Temp Title'):
        '''Writes MS Word Doc with summarized text'''
        # read or create word doc and make query  the heading

        print('Creating document')
        try:
            self.doc = docx.Document(self.outfile_path)
        except Exception as e:
            print(f'\n\nWARNING - issue creating document: {e}\n\n')
            self.doc = docx.Document()

            self.doc.add_heading(title, 1)

        try:
            self.summaries = self.do_nlp()
            self.final_text = self.clean_summaries()
            self.doc.add_paragraph(self.final_text)
            self.doc.add_heading('Extracted Tags', 2)
            self.doc.add_paragraph(', '.join(self.tags))
            self.doc.save(self.outfile_path)

        except Exception as e:
            print(f'\n\nEXCEPTION building doc: {e}\n\n')

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Scrape things!')
    parser.add_argument('infile_path', type=str, help='Word Document full filepath')
    parser.add_argument('outfile_path', type=str, help='Word Document full filepath')
    args = parser.parse_args()
    wp = Summary(args.infile_path, args.outfile_path)
    print(f'Type of wp: {type(wp)}')
    print('creating doc...')
    wp.create_text_section(args.outfile_path)
